from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from utils import normalize_space, safe_read_text


SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".md", ".txt"}


class ReportParserError(RuntimeError):
    pass


def parse_report(input_path: str | Path) -> dict[str, Any]:
    path = Path(input_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"Input report not found: {path}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ReportParserError(f"Unsupported input format '{path.suffix}'. Supported: {supported}")

    if path.suffix.lower() in {".md", ".txt"}:
        blocks = _parse_markdown_or_text(path)
    elif path.suffix.lower() == ".docx":
        blocks = _parse_docx(path)
    else:
        blocks = _parse_pdf(path)

    for idx, block in enumerate(blocks, start=1):
        block["id"] = f"b{idx:04d}"

    return {
        "source_path": str(path),
        "source_name": path.name,
        "format": path.suffix.lower().lstrip("."),
        "block_count": len(blocks),
        "blocks": blocks,
    }


def _parse_markdown_or_text(path: Path) -> list[dict[str, Any]]:
    text = safe_read_text(path)
    lines = text.splitlines()
    blocks: list[dict[str, Any]] = []
    paragraph: list[str] = []
    table_lines: list[str] = []
    list_items: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append({"type": "paragraph", "text": normalize_space(" ".join(paragraph))})
            paragraph = []

    def flush_list() -> None:
        nonlocal list_items
        if list_items:
            blocks.append({"type": "list", "items": list_items[:]})
            list_items = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            blocks.append({"type": "table", "rows": _parse_markdown_table(table_lines)})
            table_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            flush_list()
            flush_table()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            flush_list()
            flush_table()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": normalize_space(heading_match.group(2)),
                }
            )
            continue

        if "|" in stripped and stripped.count("|") >= 2:
            flush_paragraph()
            flush_list()
            table_lines.append(stripped)
            continue

        list_match = re.match(r"^\s*(?:[-*+]|\d+[.)])\s+(.+)$", line)
        if list_match:
            flush_paragraph()
            flush_table()
            list_items.append(normalize_space(list_match.group(1)))
            continue

        flush_table()
        flush_list()
        paragraph.append(stripped)

    flush_paragraph()
    flush_list()
    flush_table()

    if not blocks and text.strip():
        blocks.append({"type": "paragraph", "text": normalize_space(text)})
    return blocks


def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [normalize_space(cell) for cell in line.strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            rows.append(cells)
    return rows


def _parse_docx(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ReportParserError("python-docx is required for .docx input. Install requirements.txt.") from exc

    document = Document(str(path))
    blocks: list[dict[str, Any]] = []
    pending_list: list[str] = []

    def flush_list() -> None:
        nonlocal pending_list
        if pending_list:
            blocks.append({"type": "list", "items": pending_list[:]})
            pending_list = []

    for paragraph in document.paragraphs:
        text = normalize_space(paragraph.text)
        if not text:
            flush_list()
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            flush_list()
            level_match = re.search(r"(\d+)", style_name)
            level = int(level_match.group(1)) if level_match else 1
            blocks.append({"type": "heading", "level": level, "text": text})
        elif "list" in style_name.lower() or re.match(r"^(?:[-*+]|\d+[.)])\s+", text):
            pending_list.append(re.sub(r"^(?:[-*+]|\d+[.)])\s+", "", text))
        else:
            flush_list()
            blocks.append({"type": "paragraph", "text": text})

    flush_list()

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([normalize_space(cell.text) for cell in row.cells])
        if rows:
            blocks.append({"type": "table", "rows": rows})

    return blocks


def _parse_pdf(path: Path) -> list[dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ReportParserError("pypdf is required for .pdf input. Install requirements.txt.") from exc

    reader = PdfReader(str(path))
    blocks: list[dict[str, Any]] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_blocks = _text_to_blocks(text, page_number=page_number)
        blocks.extend(page_blocks)
    return blocks


def _text_to_blocks(text: str, page_number: int | None = None) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for part in re.split(r"\n\s*\n", text):
        clean = normalize_space(part)
        if not clean:
            continue
        block: dict[str, Any]
        if len(clean) <= 90 and not clean.endswith((".", ":", ";")):
            block = {"type": "heading", "level": 2, "text": clean}
        else:
            block = {"type": "paragraph", "text": clean}
        if page_number is not None:
            block["page"] = page_number
        blocks.append(block)
    return blocks
